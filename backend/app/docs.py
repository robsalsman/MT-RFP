"""Download attached RFP documents and extract text.

pdfplumber for digital PDFs; OCR fallback (pytesseract + pdf2image) for
scanned docs when those optional deps are installed; python-docx for the
.docx attachments some applicants upload.
"""
import logging
import re
from pathlib import Path
from urllib.parse import unquote, urlparse

import httpx

from . import config

log = logging.getLogger(__name__)

MAX_DOC_BYTES = 40 * 1024 * 1024
MAX_TEXT_CHARS = 400_000
OCR_TRIGGER_CHARS = 200  # fewer extracted chars than this => try OCR


def _safe_name(url: str) -> str:
    name = unquote(Path(urlparse(url).path).name) or "document"
    return re.sub(r"[^A-Za-z0-9._-]", "_", name)[:150]


def download_documents(application_number: str, urls: list[str]) -> list[Path]:
    """Download each doc under data/rfp_docs/{application_number}/ (skips
    files already on disk). Returns local paths."""
    dest_dir = config.DOCS_DIR / application_number
    dest_dir.mkdir(parents=True, exist_ok=True)
    paths = []
    for url in urls:
        if not url:
            continue
        dest = dest_dir / _safe_name(url)
        if dest.exists() and dest.stat().st_size > 0:
            paths.append(dest)
            continue
        try:
            with httpx.stream("GET", url, timeout=120,
                              follow_redirects=True) as resp:
                resp.raise_for_status()
                size = 0
                with open(dest, "wb") as f:
                    for chunk in resp.iter_bytes(65536):
                        size += len(chunk)
                        if size > MAX_DOC_BYTES:
                            raise ValueError("document exceeds size cap")
                        f.write(chunk)
            paths.append(dest)
        except Exception as e:  # network hiccups shouldn't sink the sync
            log.warning("download failed %s: %s", url, e)
            dest.unlink(missing_ok=True)
    return paths


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf(path)[:MAX_TEXT_CHARS]
        if suffix == ".docx":
            return _extract_docx(path)[:MAX_TEXT_CHARS]
        if suffix in (".txt", ".csv"):
            return path.read_text(errors="replace")[:MAX_TEXT_CHARS]
    except Exception as e:
        log.warning("text extraction failed %s: %s", path.name, e)
    return ""


def _extract_pdf(path: Path) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    text = "\n".join(text_parts).strip()
    if len(text) < OCR_TRIGGER_CHARS:
        ocr = _ocr_pdf(path)
        if len(ocr) > len(text):
            return ocr
    return text


def _ocr_pdf(path: Path) -> str:
    """Best-effort OCR; requires pytesseract + pdf2image + poppler/tesseract
    binaries. Silently returns '' when unavailable."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError:
        log.info("OCR deps not installed; skipping OCR for %s", path.name)
        return ""
    try:
        pages = convert_from_path(str(path), dpi=200, last_page=40)
        return "\n".join(pytesseract.image_to_string(p) for p in pages).strip()
    except Exception as e:
        log.warning("OCR failed %s: %s", path.name, e)
        return ""


def _extract_docx(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts).strip()
