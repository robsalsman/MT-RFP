"""Closing documents: the artifacts that get deals signed.

- Savings sheet: one page, their real numbers vs Mission pricing — the
  thing a tech director forwards to a superintendent.
- Champion kit: the internal pitch (board memo) Kim's contact uses to
  sell the switch in rooms Kim isn't in.
- Case study: post-win, fuels the next three deals.

All facts come from the lead row (real USAC/ECF data). Estimates are
labeled as estimates. AI writes connective prose when available; a
deterministic template takes over offline — documents always generate.
"""
import datetime
import logging
import re
from pathlib import Path

from . import ai, competitors, config, db

log = logging.getLogger(__name__)

OUT_DIR = Path(config.RESPONSES_DIR).parent / "closing"

# Mission hotspot/broadband pricing (missiontelecom.org)
MISSION_LOW, MISSION_HIGH = 20.0, 25.0
# typical K-12 market rate per managed hotspot line (used ONLY when the
# lead's own device count is unknown; always labeled an estimate)
MARKET_LOW, MARKET_HIGH = 35.0, 40.0


def erate_sign_by() -> dict:
    """E-Rate timing: to be funded for the next funding year, competitive
    bidding (28-day Form 470) must finish before the Form 471 filing
    window closes (typically late March)."""
    today = datetime.date.today()
    fy = today.year if today.month >= 7 else today.year - 1
    next_fy = fy + 1
    window_close = datetime.date(next_fy, 3, 25)   # typical close
    post_470_by = window_close - datetime.timedelta(days=45)
    return {"next_funding_year": next_fy,
            "service_start": f"July {next_fy}",
            "form_471_window_closes": window_close.isoformat(),
            "post_470_by": post_470_by.isoformat()}


def compute_savings(lead: dict) -> dict:
    """Savings math from real numbers; honest fallbacks when device count
    is unknown."""
    spend = float(lead.get("spend") or 0)
    devices = lead.get("devices")
    yearly = spend if lead.get("source") != "ecf" else None
    out = {"current_annual": yearly, "ecf_total": spend if yearly is None
           else None, "devices": devices, "estimated": False}
    if devices and yearly:
        per_line = yearly / 12 / devices
        mission_lo, mission_hi = devices * 12 * MISSION_LOW, \
            devices * 12 * MISSION_HIGH
        out.update({"current_per_line": round(per_line, 2),
                    "mission_annual_low": round(mission_lo, 2),
                    "mission_annual_high": round(mission_hi, 2),
                    "savings_low": round(yearly - mission_hi, 2),
                    "savings_high": round(yearly - mission_lo, 2)})
    elif devices:
        # ECF lead: device count known, annual spend is program-total
        out.update({"estimated": False,
                    "mission_annual_low": round(devices * 12 * MISSION_LOW, 2),
                    "mission_annual_high": round(devices * 12 * MISSION_HIGH, 2)})
    elif yearly:
        # no device count: bracket with market-rate assumption, labeled
        out.update({"estimated": True,
                    "savings_low": round(yearly * (1 - MISSION_HIGH / MARKET_LOW), 2),
                    "savings_high": round(yearly * (1 - MISSION_LOW / MARKET_HIGH), 2)})
    return out


def _doc(title: str):
    import docx
    from docx.shared import Pt, RGBColor
    d = docx.Document()
    h = d.add_heading(title, level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x5F, 0xFF)
    p = d.add_paragraph("Mission Telecom — nonprofit wireless broadband on "
                        "the T-Mobile network  ·  missiontelecom.org  ·  "
                        "877-641-9444")
    p.runs[0].font.size = Pt(9)
    return d


def _fact_lines(lead: dict, sv: dict) -> list[str]:
    comp = lead.get("competitor_label") or "the incumbent"
    lines = [f"Organization: {lead['org']} ({lead['state']})"]
    if sv["current_annual"]:
        lines.append(f"Current mobile-broadband spend with {comp}: "
                     f"${sv['current_annual']:,.0f}/year "
                     "(public E-Rate Form 471 data)")
    if sv.get("ecf_total"):
        lines.append(f"ECF-funded hotspot program with {comp}: "
                     f"${sv['ecf_total']:,.0f} (program has ended)")
    if sv.get("devices"):
        lines.append(f"Hotspot lines on record: {sv['devices']:,}")
    if sv.get("current_per_line"):
        lines.append(f"Effective cost per line: "
                     f"${sv['current_per_line']:.2f}/month")
    if lead.get("next_expiration"):
        lines.append(f"Current contract runs through "
                     f"{lead['next_expiration']}")
    if lead.get("enrollment"):
        lines.append(f"Students served: ~{lead['enrollment']:,}")
    return lines


def _savings_block(d, sv):
    import docx  # noqa: F401
    d.add_heading("The numbers", level=1)
    t = d.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"

    def row(a, b):
        cells = t.add_row().cells
        cells[0].text = a
        cells[1].text = b
    if sv.get("current_annual"):
        row("Current annual spend", f"${sv['current_annual']:,.0f}")
    if sv.get("current_per_line"):
        row("Current cost per line", f"${sv['current_per_line']:.2f}/mo")
    row("Mission Telecom rate", f"${MISSION_LOW:.0f}–{MISSION_HIGH:.0f}"
        "/line/mo (nonprofit pricing, free CIPA-compliant filtering)")
    if sv.get("mission_annual_low"):
        row("Mission annual cost (same line count)",
            f"${sv['mission_annual_low']:,.0f}–"
            f"${sv['mission_annual_high']:,.0f}")
    if sv.get("savings_low") is not None:
        label = "Estimated annual savings" if sv["estimated"] \
            else "Annual savings"
        row(label, f"${max(sv['savings_low'], 0):,.0f}–"
            f"${max(sv['savings_high'], 0):,.0f}")
    if sv["estimated"]:
        p = d.add_paragraph(
            "Savings estimated against typical K-12 managed-hotspot market "
            f"rates (${MARKET_LOW:.0f}–{MARKET_HIGH:.0f}/line/mo); we'll "
            "firm this up against your actual line count.")
        p.runs[0].italic = True


def _timing_block(d):
    cal = erate_sign_by()
    d.add_heading("Timing (E-Rate)", level=1)
    d.add_paragraph(
        f"To fund this for FY{cal['next_funding_year']} (service starting "
        f"{cal['service_start']}): post the FCC Form 470 by "
        f"{cal['post_470_by']} (28-day competitive bidding), then file the "
        f"Form 471 before the window closes (~"
        f"{cal['form_471_window_closes']}). Mission Telecom's nonprofit "
        "pricing also works without E-Rate — you can start immediately and "
        "add funding next cycle.")


def _ai_intro(kind: str, lead: dict, sv: dict) -> str | None:
    facts = "; ".join(_fact_lines(lead, sv))
    return ai._chat(
        "Write ONE short paragraph (60-90 words), plain text, no heading, "
        f"for the opening of a {kind} aimed at a school/library decision-"
        "maker. Professional, warm, concrete. Use ONLY the facts given. "
        "Mission Telecom is a nonprofit wireless ISP on the T-Mobile "
        "network offering hotspot lending and mobile broadband with free "
        "CIPA-compliant filtering.", f"Facts: {facts}", max_tokens=800)


def build_doc(lead_id: int, kind: str) -> dict:
    """kind: savings | champion | case. Returns {path, filename}."""
    lead = competitors.get_lead(lead_id)
    if not lead:
        return {"error": "no such lead"}
    sv = compute_savings(lead)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    safe = re.sub(r"[^A-Za-z0-9]+", "_", lead["org"])[:40].strip("_")
    comp = lead.get("competitor_label") or "incumbent"

    if kind == "savings":
        d = _doc(f"Connectivity savings for {lead['org']}")
        intro = _ai_intro("savings summary", lead, sv) or (
            f"{lead['org']} currently funds mobile broadband through "
            f"{comp}. As a nonprofit carrier on the T-Mobile network, "
            "Mission Telecom delivers the same connectivity at nonprofit "
            "rates — the summary below uses your own public E-Rate filing "
            "data.")
        d.add_paragraph(intro)
        d.add_heading("Where you are today", level=1)
        for line in _fact_lines(lead, sv):
            d.add_paragraph(line, style="List Bullet")
        _savings_block(d, sv)
        _timing_block(d)
        d.add_heading("Next step", level=1)
        d.add_paragraph("A 15-minute call to confirm line counts and "
                        "coverage, then a firm quote within 48 hours.")
        name = f"Savings_{safe}.docx"

    elif kind == "champion":
        d = _doc(f"Board briefing: student & patron connectivity — "
                 f"{lead['org']}")
        intro = _ai_intro("internal board briefing (written for staff to "
                          "present to leadership)", lead, sv) or (
            "This briefing summarizes an opportunity to reduce our "
            "connectivity costs and expand hotspot access by moving to "
            "Mission Telecom, a nonprofit carrier on the T-Mobile network.")
        d.add_paragraph(intro)
        d.add_heading("Current situation", level=1)
        for line in _fact_lines(lead, sv):
            d.add_paragraph(line, style="List Bullet")
        try:
            from . import acp
            hh = acp.households_for_zip(lead.get("zip"))
            if hh and hh > 200:
                d.add_paragraph(
                    f"Community need: ~{hh:,} households in our zip code "
                    "lost the federal ACP internet subsidy when the "
                    "program ended in 2024.", style="List Bullet")
        except Exception:
            pass
        _savings_block(d, sv)
        d.add_heading("Compliance & eligibility", level=1)
        for x in ("E-Rate eligible (Category 1 data transmission / "
                  "internet access)",
                  "CIPA-compliant content filtering included at no cost",
                  "Nonprofit carrier — mission-aligned pricing, "
                  "5-year price guarantees available",
                  "Runs on the nationwide T-Mobile 5G/4G network"):
            d.add_paragraph(x, style="List Bullet")
        _timing_block(d)
        d.add_heading("Recommendation", level=1)
        d.add_paragraph("Authorize a no-cost pilot and a firm quote from "
                        "Mission Telecom for side-by-side comparison at "
                        "the next renewal decision.")
        name = f"Board_Briefing_{safe}.docx"

    elif kind == "case":
        d = _doc(f"Case study: {lead['org']}")
        body = ai._chat(
            "Write a 150-200 word draft case study (plain text, no "
            "headings) about a school/library that switched to Mission "
            "Telecom (nonprofit wireless ISP, T-Mobile network, hotspot "
            "lending, free CIPA filtering). Use ONLY the facts given; "
            "where outcomes are unknown, leave [NEEDS INPUT] markers "
            "rather than inventing results.",
            "Facts: " + "; ".join(_fact_lines(lead, sv)), max_tokens=900) \
            or (f"{lead['org']} ({lead['state']}) moved its mobile-"
                f"broadband service from {comp} to Mission Telecom. "
                "[NEEDS INPUT: deployment size, timeline, outcomes, "
                "quote from the customer.]")
        d.add_paragraph(body)
        d.add_paragraph("DRAFT — verify all facts and obtain customer "
                        "approval before any external use.").runs[0].bold \
            = True
        name = f"Case_Study_{safe}.docx"
    else:
        return {"error": f"unknown document kind '{kind}'"}

    path = OUT_DIR / name
    d.save(str(path))
    return {"path": str(path), "filename": name, "kind": kind,
            "savings": sv}
