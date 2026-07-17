"""One-click draft response generation: cover letter, compliance matrix,
pricing table (deterministic, from the uploaded price list only), required-
forms checklist, exported as DOCX + PDF. Every output is a DRAFT with a
human-review checklist; nothing is ever submitted anywhere.
"""
import json
import re
from datetime import datetime, timezone

from . import ai, config, db, pricing

NEEDS_INPUT = "[NEEDS INPUT]"

REVIEW_CHECKLIST = [
    "Verify every [NEEDS INPUT] placeholder has been filled by a human",
    "Confirm pricing table quantities and totals against the RFP line items",
    "Red-flagged (unmatched) pricing rows priced manually",
    "Compliance matrix reviewed by someone with bid authority",
    "Signatures / notarization obtained where required",
    "Bid bond / surety attached if the RFP requires one",
    "Insurance certificates attached (from company document uploads)",
    "W-9 attached and current",
    "Submission method and deadline double-checked (applicant-local AND Eastern)",
    "Final proofread — remove this checklist page before submitting",
]


def generate_response(application_number: str) -> dict:
    with db.closing_conn() as conn:
        row = conn.execute("SELECT * FROM rfps WHERE application_number=?",
                           (application_number,)).fetchone()
        srs = [dict(r) for r in conn.execute(
            "SELECT * FROM service_requests WHERE application_number=?",
            (application_number,)).fetchall()]
        profile = db.kv_get(conn, "company_profile", {})
    if row is None:
        raise ValueError(f"unknown application {application_number}")
    row = dict(row)

    analysis = None
    if row.get("analysis"):
        analysis = json.loads(row["analysis"])
    if analysis is None:
        analysis = ai.analyze_rfp(application_number) or _fallback_analysis(row)

    narratives = ai.draft_narratives(row, analysis, profile)
    _enforce_no_fabrication(narratives, profile)
    price_rows = pricing.match_services(srs)
    unmatched = [p for p in price_rows if not p["matched"]]

    entity = re.sub(r"[^A-Za-z0-9]+", "_",
                    row.get("billed_entity_name") or "Entity").strip("_")[:60]
    base = f"{entity}_{application_number}_MissionTelecom_Response"
    docx_path = config.RESPONSES_DIR / f"{base}.docx"
    pdf_path = config.RESPONSES_DIR / f"{base}.pdf"

    _write_docx(docx_path, row, analysis, narratives, price_rows, profile)
    _write_pdf(pdf_path, row, analysis, narratives, price_rows, profile)

    with db.closing_conn() as conn:
        cur = conn.execute(
            "INSERT INTO responses (application_number, created_at, docx_path,"
            " pdf_path, checklist, unmatched_items, status) "
            "VALUES (?,?,?,?,?,?, 'DRAFT')",
            (application_number, datetime.now(timezone.utc).isoformat(),
             str(docx_path), str(pdf_path), json.dumps(REVIEW_CHECKLIST),
             json.dumps([u["request"] for u in unmatched])))
        conn.commit()
        response_id = cur.lastrowid
    return {"id": response_id, "docx": str(docx_path), "pdf": str(pdf_path),
            "unmatched_count": len(unmatched),
            "checklist": REVIEW_CHECKLIST, "status": "DRAFT"}


def _fallback_analysis(row: dict) -> dict:
    """Minimal analysis from structured 470 data when no AI key is set."""
    return {
        "issuing_entity": row.get("billed_entity_name"),
        "state": row.get("state"),
        "services_requested": [
            {"service": s, "quantity": None, "bandwidth": None}
            for s in json.loads(row.get("service_types") or "[]")],
        "contract_term_years": None,
        "mandatory_requirements": [],
        "evaluation_criteria": [],
        "price_primary_factor": None,
        "question_deadline": None,
        "submission_method": None,
        "submission_deadline": None,
        "disqualifiers": [],
        "rationale": row.get("score_rationale") or "",
    }


def _enforce_no_fabrication(narratives: dict, profile: dict) -> None:
    """Belt-and-braces guardrail on top of the model prompt: strip dollar
    amounts from narratives (prices only come from the price list) and blank
    company identifiers that aren't in the uploaded profile."""
    profile_blob = json.dumps(profile).lower()
    def scrub(text: str) -> str:
        # no model-authored dollar figures, ever
        text = re.sub(r"\$\s?[\d,]+(?:\.\d+)?", NEEDS_INPUT, text or "")
        # SPIN/FCC RN style identifiers must exist in the profile verbatim
        for ident in re.findall(r"\b\d{9,10}\b", text):
            if ident.lower() not in profile_blob:
                text = text.replace(ident, NEEDS_INPUT)
        return text
    narratives["cover_letter"] = scrub(narratives.get("cover_letter", ""))
    for item in narratives.get("compliance", []):
        item["response"] = scrub(item.get("response", ""))


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------

def _write_docx(path, row, analysis, narratives, price_rows, profile):
    import docx
    from docx.shared import Pt, RGBColor

    d = docx.Document()
    RED = RGBColor(0xC0, 0x00, 0x00)

    d.add_heading("DRAFT — NOT FOR SUBMISSION", 0)
    p = d.add_paragraph()
    run = p.add_run("Machine-generated draft. Human review is mandatory before"
                    " any use. See review checklist on the final page.")
    run.bold = True
    run.font.color.rgb = RED

    d.add_heading(f"Response to FCC Form 470 #{row['application_number']}", 1)
    meta = d.add_paragraph()
    meta.add_run(f"{row.get('billed_entity_name')} — {row.get('city')}, "
                 f"{row.get('state')}\n")
    meta.add_run(f"Funding Year {row.get('funding_year')} | Certified "
                 f"{(row.get('certified_date') or '')[:10]} | Allowable "
                 f"Contract Date {(row.get('allowable_contract_date') or '')[:10]}")

    d.add_heading("1. Cover Letter", 1)
    for para in (narratives.get("cover_letter") or NEEDS_INPUT).split("\n\n"):
        d.add_paragraph(para)

    d.add_heading("2. Compliance Matrix", 1)
    comp = narratives.get("compliance") or []
    if comp:
        t = d.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = (
            "RFP Requirement", "Mission Telecom Response", "Status")
        for item in comp:
            c = t.add_row().cells
            c[0].text = item.get("requirement", "")
            c[1].text = item.get("response", NEEDS_INPUT)
            c[2].text = item.get("compliance", "REVIEW")
    else:
        d.add_paragraph("No mandatory requirements were extracted from the "
                        "RFP documents. " + NEEDS_INPUT +
                        ": review the original RFP for requirements.")

    d.add_heading("3. Pricing Table", 1)
    d.add_paragraph("Prices come exclusively from the uploaded Mission "
                    "Telecom price list. Rows in red had no price-list match "
                    "and require human pricing.")
    t = d.add_table(rows=1, cols=7)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Requested Service", "Capacity", "Qty", "SKU",
                           "Unit Price", "Extended", "Match"]):
        t.rows[0].cells[i].text = h
    for pr in price_rows:
        sr = pr["request"]
        c = t.add_row().cells
        c[0].text = " / ".join(x for x in (sr.get("service_type"),
                                           sr.get("function")) if x)
        c[1].text = sr.get("max_capacity") or sr.get("min_capacity") or ""
        c[2].text = str(pr.get("quantity") or sr.get("quantity") or "")
        if pr["matched"]:
            c[3].text = pr["sku"] or ""
            c[4].text = f"${pr['unit_price']:,.2f}"
            c[5].text = (f"${pr['extended_price']:,.2f}"
                         if pr["extended_price"] is not None else NEEDS_INPUT)
            c[6].text = pr["confidence"]
        else:
            for idx, txt in ((3, NEEDS_INPUT), (4, NEEDS_INPUT),
                             (5, NEEDS_INPUT), (6, "NO MATCH")):
                para = c[idx].paragraphs[0]
                run = para.add_run(txt)
                run.font.color.rgb = RED
                run.bold = True
    total = sum(p["extended_price"] or 0 for p in price_rows if p["matched"])
    tp = d.add_paragraph()
    tp.add_run(f"Subtotal of matched items: ${total:,.2f} ").bold = True
    tp.add_run("(excludes unmatched rows — incomplete until all rows are "
               "priced)")

    d.add_heading("4. Company Information", 1)
    fields = [("Legal name", "legal_name"), ("SPIN #", "spin"),
              ("FCC RN", "fcc_rn"), ("Address", "address"),
              ("Contact", "contact_name"), ("Contact email", "contact_email"),
              ("Contact phone", "contact_phone")]
    for label, key in fields:
        d.add_paragraph(f"{label}: {profile.get(key) or NEEDS_INPUT}")
    if profile.get("capability_statement"):
        d.add_heading("Capability Statement", 2)
        d.add_paragraph(profile["capability_statement"])
    refs = profile.get("references") or []
    d.add_heading("References", 2)
    if refs:
        for r in refs:
            d.add_paragraph(str(r), style="List Bullet")
    else:
        d.add_paragraph(NEEDS_INPUT + ": add standard references in company "
                        "profile")

    d.add_heading("5. Required Forms & Human Actions Checklist", 1)
    for item in REVIEW_CHECKLIST:
        d.add_paragraph(item, style="List Bullet")

    for section in d.sections:
        footer_p = section.footer.paragraphs[0]
        footer_p.text = ("DRAFT — Mission Telecom — generated "
                         f"{datetime.now(timezone.utc).strftime('%Y-%m-%d')}"
                         " — human review required")
        footer_p.runs[0].font.size = Pt(8)
    d.save(str(path))


# --------------------------------------------------------------------------
# PDF (reportlab rendering of the same content)
# --------------------------------------------------------------------------

def _write_pdf(path, row, analysis, narratives, price_rows, profile):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer,
                                    Table, TableStyle)

    styles = getSampleStyleSheet()
    story = [
        Paragraph("DRAFT — NOT FOR SUBMISSION", styles["Title"]),
        Paragraph(f"Response to FCC Form 470 #{row['application_number']} — "
                  f"{row.get('billed_entity_name')} ({row.get('state')})",
                  styles["Heading1"]),
        Spacer(1, 12),
        Paragraph("1. Cover Letter", styles["Heading2"]),
    ]
    for para in (narratives.get("cover_letter") or NEEDS_INPUT).split("\n\n"):
        story.append(Paragraph(para.replace("\n", "<br/>"), styles["Normal"]))
        story.append(Spacer(1, 6))

    story.append(Paragraph("2. Compliance Matrix", styles["Heading2"]))
    comp = narratives.get("compliance") or []
    if comp:
        data = [["Requirement", "Response", "Status"]] + [
            [Paragraph(i.get("requirement", ""), styles["BodyText"]),
             Paragraph(i.get("response", NEEDS_INPUT), styles["BodyText"]),
             i.get("compliance", "REVIEW")] for i in comp]
        t = Table(data, colWidths=[2.4 * inch, 3.2 * inch, 0.9 * inch])
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey)]))
        story.append(t)
    else:
        story.append(Paragraph("No mandatory requirements extracted. "
                               + NEEDS_INPUT, styles["Normal"]))

    story.append(Spacer(1, 10))
    story.append(Paragraph("3. Pricing Table", styles["Heading2"]))
    data = [["Requested Service", "Capacity", "Qty", "SKU", "Unit Price",
             "Extended"]]
    red_rows = []
    for i, pr in enumerate(price_rows, start=1):
        sr = pr["request"]
        svc = " / ".join(x for x in (sr.get("service_type"),
                                     sr.get("function")) if x)
        if pr["matched"]:
            data.append([Paragraph(svc, styles["BodyText"]),
                         sr.get("max_capacity") or "",
                         str(pr.get("quantity") or ""), pr["sku"] or "",
                         f"${pr['unit_price']:,.2f}",
                         (f"${pr['extended_price']:,.2f}"
                          if pr["extended_price"] is not None else NEEDS_INPUT)])
        else:
            red_rows.append(i)
            data.append([Paragraph(svc, styles["BodyText"]),
                         sr.get("max_capacity") or "",
                         str(pr.get("quantity") or ""),
                         NEEDS_INPUT, NEEDS_INPUT, NEEDS_INPUT])
    t = Table(data, colWidths=[2.4 * inch, 0.9 * inch, 0.5 * inch,
                               1.0 * inch, 0.9 * inch, 0.9 * inch])
    style = [("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
             ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey)]
    for ri in red_rows:
        style.append(("TEXTCOLOR", (0, ri), (-1, ri), colors.red))
    t.setStyle(TableStyle(style))
    story.append(t)

    story.append(Spacer(1, 10))
    story.append(Paragraph("4. Company Information", styles["Heading2"]))
    for label, key in [("Legal name", "legal_name"), ("SPIN #", "spin"),
                       ("FCC RN", "fcc_rn"), ("Address", "address")]:
        story.append(Paragraph(f"{label}: {profile.get(key) or NEEDS_INPUT}",
                               styles["Normal"]))

    story.append(Spacer(1, 10))
    story.append(Paragraph("5. Required Forms & Human Actions Checklist",
                           styles["Heading2"]))
    for item in REVIEW_CHECKLIST:
        story.append(Paragraph(f"• {item}", styles["Normal"]))

    SimpleDocTemplate(str(path), pagesize=letter).build(story)
